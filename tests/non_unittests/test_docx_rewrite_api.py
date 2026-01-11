import os
import sys
import uuid
import urllib.request
from io import BytesIO


ROOT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
if ROOT_DIR not in sys.path:
    sys.path.insert(0, ROOT_DIR)


def _read_pdf_text(pdf_path: str, max_chars: int = 3000) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return ""
    try:
        reader = PdfReader(pdf_path)
    except Exception:
        return ""
    parts = []
    total = 0
    for page in getattr(reader, "pages", []) or []:
        try:
            t = (page.extract_text() or "").strip()
        except Exception:
            t = ""
        if not t:
            continue
        if total + len(t) > max_chars:
            remain = max_chars - total
            if remain <= 0:
                break
            t = t[:remain]
        parts.append(t)
        total += len(t)
        if total >= max_chars:
            break
    return "\n\n".join(parts).strip()


def _encode_multipart_formdata(fields: dict, files: dict) -> tuple[bytes, str]:
    boundary = f"----agenticrag{uuid.uuid4().hex}"
    lines: list[bytes] = []

    for name, value in (fields or {}).items():
        lines.append(f"--{boundary}".encode("utf-8"))
        lines.append(f'Content-Disposition: form-data; name="{name}"'.encode("utf-8"))
        lines.append(b"")
        lines.append(str(value).encode("utf-8"))

    for name, fileinfo in (files or {}).items():
        filename = fileinfo["filename"]
        content = fileinfo["content"]
        content_type = fileinfo.get("content_type") or "application/octet-stream"
        lines.append(f"--{boundary}".encode("utf-8"))
        lines.append(
            (
                f'Content-Disposition: form-data; name="{name}"; filename="{filename}"\r\n'
                f"Content-Type: {content_type}"
            ).encode("utf-8")
        )
        lines.append(b"")
        lines.append(content)

    lines.append(f"--{boundary}--".encode("utf-8"))
    lines.append(b"")
    body = b"\r\n".join(lines)
    return body, f"multipart/form-data; boundary={boundary}"


def _post_multipart(url: str, fields: dict, files: dict, headers: dict | None = None) -> tuple[int, bytes, dict]:
    body, content_type = _encode_multipart_formdata(fields, files)
    req_headers = {"Content-Type": content_type}
    if headers:
        req_headers.update(headers)
    req = urllib.request.Request(url, data=body, headers=req_headers, method="POST")
    try:
        with urllib.request.urlopen(req) as resp:
            return resp.status, resp.read(), dict(resp.headers)
    except urllib.error.HTTPError as e:
        try:
            return int(getattr(e, "code", 0) or 0), e.read(), dict(getattr(e, "headers", {}) or {})
        except Exception:
            return int(getattr(e, "code", 0) or 0), str(e).encode("utf-8", errors="ignore"), {}
    except Exception as e:
        msg = str(e).encode("utf-8", errors="ignore")
        return 0, msg, {}


def test_docx_rewrite_endpoint_with_template_and_pdf_content():
    from dotenv import load_dotenv

    load_dotenv()
    api_key = (os.getenv("DEEPSEEK_API_KEY") or "").strip() or (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        assert True
        return

    try:
        from docx import Document  # type: ignore
    except Exception:
        assert True
        return

    template_path = os.path.join(ROOT_DIR, "tests", "testfiles", "330 Wu Ka Long, Long.docx")
    pdf_path = os.path.join(ROOT_DIR, "tests", "testfiles", "严冬-个人简历2040721.docx.pdf")
    assert os.path.isfile(template_path)
    assert os.path.isfile(pdf_path)

    pdf_text = _read_pdf_text(pdf_path, max_chars=3500)
    instruction = (
        "请参考模板文档的排版与写作风格，基于以下简历内容，重新生成一份中文个人简历。"
        "要求：信息完整、结构清晰、语言专业、避免虚构。\n\n"
        "简历内容（从PDF提取，可能不完整）：\n"
        f"{pdf_text}"
    )

    with open(template_path, "rb") as f:
        tpl_bytes = f.read()
    tpl_doc = Document(template_path)
    tpl_para_count = len(tpl_doc.paragraphs)

    headers = {
        "x-llm-api-key": api_key,
        "x-llm-base-url": (os.getenv("LLM_BASE_URL") or "").strip(),
        "x-llm-model": (os.getenv("LLM_MODEL") or "").strip(),
    }
    status, out_bytes, resp_headers = _post_multipart(
        "http://127.0.0.1:8000/api/docx/rewrite",
        fields={"instruction_text": instruction},
        files={
            "reference_file": {
                "filename": os.path.basename(template_path),
                "content": tpl_bytes,
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }
        },
        headers={k: v for k, v in headers.items() if v},
    )
    assert status == 200, f"unexpected status: {status}, body: {out_bytes[:500]!r}"
    assert out_bytes[:2] == b"PK", "response is not a docx(zip)"

    out_dir = os.path.join(ROOT_DIR, "tests", "outputs", "docx")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"rewrite_output_{uuid.uuid4().hex}.docx")
    with open(out_path, "wb") as f:
        f.write(out_bytes)
    print(f"docx rewrite output saved: {out_path}")

    out_doc = Document(BytesIO(out_bytes))
    assert len(out_doc.paragraphs) == tpl_para_count
    assert any((p.text or "").strip() for p in out_doc.paragraphs)
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    content_type = ""
    for k, v in (resp_headers or {}).items():
        if str(k).lower() == "content-type":
            content_type = str(v)
            break
    if content_type:
        assert docx_mime in content_type


def main():
    test_docx_rewrite_endpoint_with_template_and_pdf_content()


if __name__ == "__main__":
    main()

