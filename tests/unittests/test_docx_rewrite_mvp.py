import unittest
from io import BytesIO


class TestDocxRewriteMvp(unittest.TestCase):
    def test_extract_and_render_keeps_paragraph_count(self):
        from docx import Document  # type: ignore

        from backend.services.docx_rewrite_service import extract_docx_template, render_rewritten_docx

        doc = Document()
        doc.add_heading("原始标题", level=1)
        doc.add_paragraph("第一段正文。")
        doc.add_paragraph("")
        doc.add_paragraph("第二段正文。")

        buf = BytesIO()
        doc.save(buf)
        src = buf.getvalue()

        template = extract_docx_template(src)
        self.assertEqual(len(template), 4)

        mapping = {
            0: "新标题",
            1: "替换后的第一段正文。",
            3: "替换后的第二段正文。",
        }
        out = render_rewritten_docx(src, mapping)
        out_doc = Document(BytesIO(out))
        texts = [p.text for p in out_doc.paragraphs]

        self.assertEqual(len(texts), 4)
        self.assertEqual(texts[0], "新标题")
        self.assertEqual(texts[1], "替换后的第一段正文。")
        self.assertEqual(texts[2], "")
        self.assertEqual(texts[3], "替换后的第二段正文。")

    def test_render_preserves_first_run_formatting(self):
        from docx import Document  # type: ignore

        from backend.services.docx_rewrite_service import render_rewritten_docx

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("BOLD")
        r.bold = True
        r.font.name = "Calibri"

        buf = BytesIO()
        doc.save(buf)
        src = buf.getvalue()

        out = render_rewritten_docx(src, {0: "NEW TEXT"})
        out_doc = Document(BytesIO(out))
        out_runs = out_doc.paragraphs[0].runs
        self.assertTrue(len(out_runs) >= 1)
        self.assertEqual(out_doc.paragraphs[0].text, "NEW TEXT")
        self.assertEqual(out_runs[0].bold, True)

    def test_render_enforces_case_from_template_terms(self):
        from docx import Document  # type: ignore

        from backend.services.docx_rewrite_service import render_rewritten_docx

        doc = Document()
        doc.add_paragraph("SKILLS: PYTHON")

        buf = BytesIO()
        doc.save(buf)
        src = buf.getvalue()

        out = render_rewritten_docx(src, {0: "skills: python"})
        out_doc = Document(BytesIO(out))
        self.assertEqual(out_doc.paragraphs[0].text, "SKILLS: PYTHON")

    def test_tables_cells_are_extracted_and_rewritten_with_format(self):
        from docx import Document  # type: ignore

        from backend.services.docx_rewrite_service import extract_docx_template, render_rewritten_docx

        doc = Document()
        doc.add_paragraph("Intro")
        tbl = doc.add_table(rows=1, cols=2)
        cell_l = tbl.cell(0, 0)
        p_l = cell_l.paragraphs[0]
        r_l = p_l.add_run("CPU")
        r_l.bold = True
        cell_r = tbl.cell(0, 1)
        p_r = cell_r.paragraphs[0]
        p_r.add_run("email: EXAMPLE@EX.COM")
        doc.add_paragraph("Footer")

        buf = BytesIO()
        doc.save(buf)
        src = buf.getvalue()

        template = extract_docx_template(src)
        self.assertTrue(len(template) >= 4)

        mapping = {
            0: "Intro new",
            1: "cpu",
            2: "email: example@ex.com",
            3: "Footer new",
        }
        out = render_rewritten_docx(src, mapping)
        out_doc = Document(BytesIO(out))

        self.assertEqual(out_doc.paragraphs[0].text, "Intro new")
        self.assertEqual(out_doc.paragraphs[-1].text, "Footer new")
        self.assertEqual(out_doc.tables[0].cell(0, 0).paragraphs[0].text, "CPU")
        self.assertTrue(out_doc.tables[0].cell(0, 0).paragraphs[0].runs[0].bold)
        self.assertEqual(out_doc.tables[0].cell(0, 1).paragraphs[0].text, "email: EXAMPLE@EX.COM")


if __name__ == "__main__":
    unittest.main()
