from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
import json
import os
import re
from typing import Any, Dict, Iterable, List, Optional

from dotenv import load_dotenv
from langchain_openai import ChatOpenAI


@dataclass(frozen=True)
class LlmConfig:
    api_key: str
    base_url: str
    model: str


@dataclass(frozen=True)
class TemplateParagraph:
    index: int
    style: str
    kind: str
    original_text: str
    target_len: int
    min_len: int
    max_len: int


def resolve_llm_config(
    *,
    api_key: str | None = None,
    base_url: str | None = None,
    model: str | None = None,
) -> Optional[LlmConfig]:
    load_dotenv()
    key = (
        (api_key or "").strip()
        or (os.getenv("DEEPSEEK_API_KEY") or "").strip()
        or (os.getenv("OPENAI_API_KEY") or "").strip()
    )
    if not key:
        return None
    url = (base_url or "").strip() or (os.getenv("LLM_BASE_URL") or "").strip() or "https://api.deepseek.com/v1"
    m = (model or "").strip() or (os.getenv("LLM_MODEL") or "").strip() or "deepseek-chat"
    return LlmConfig(api_key=key, base_url=url, model=m)


def estimate_text_length(text: str) -> int:
    s = (text or "").strip()
    if not s:
        return 0
    cjk = sum(1 for ch in s if "\u4e00" <= ch <= "\u9fff")
    ascii_words = len(re.findall(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?", s))
    other = len(re.sub(r"\s+", "", s)) - cjk
    return int(cjk + ascii_words + max(0, other - ascii_words))


def detect_language_hint(text: str) -> str:
    s = (text or "").strip()
    if not s:
        return "unknown"
    cjk = sum(1 for ch in s if "\u4e00" <= ch <= "\u9fff")
    latin = len(re.findall(r"[A-Za-z]", s))
    if cjk >= 20 and cjk > int(latin * 1.2):
        return "zh"
    if latin >= 20 and latin > int(cjk * 1.2):
        return "en"
    if cjk > 0 and latin > 0:
        return "mixed"
    return "zh" if cjk > 0 else "en"


def _paragraph_kind(style_name: str, text: str) -> str:
    s = (style_name or "").strip().lower()
    t = (text or "").strip()
    if not t:
        return "blank"
    if "heading" in s or "title" in s or "标题" in (style_name or ""):
        return "heading"
    if t.startswith(("- ", "* ", "• ")) or re.match(r"^\d+[\.)]\s+", t):
        return "list"
    return "body"


def extract_docx_template(reference_docx_bytes: bytes, *, max_paragraphs: int = 400) -> List[TemplateParagraph]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("缺少依赖：python-docx") from e

    doc = Document(BytesIO(reference_docx_bytes))
    body = getattr(getattr(doc, "_element", None), "body", None)
    if body is None:
        paras = list(getattr(doc, "paragraphs", []) or [])
        if len(paras) > max_paragraphs:
            raise ValueError(f"文档段落过多（{len(paras)}），请先精简后再仿写")
        out: List[TemplateParagraph] = []
        for idx, p in enumerate(paras):
            style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
            original_text = str(getattr(p, "text", "") or "")
            kind = _paragraph_kind(style_name, original_text)
            target_len = estimate_text_length(original_text)
            if kind == "blank":
                min_len = 0
                max_len = 0
            elif kind == "heading":
                min_len = max(4, int(target_len * 0.5))
                max_len = max(12, int(target_len * 1.2))
            else:
                min_len = max(20, int(target_len * 0.7))
                max_len = max(60, int(target_len * 1.3))
            out.append(
                TemplateParagraph(
                    index=idx,
                    style=style_name,
                    kind=kind,
                    original_text=original_text,
                    target_len=target_len,
                    min_len=min_len,
                    max_len=max_len,
                )
            )
        return out

    out: List[TemplateParagraph] = []
    idx = 0
    for child in list(body):
        tag = str(getattr(child, "tag", ""))
        if tag.endswith("}p"):
            from docx.text.paragraph import Paragraph  # type: ignore
            p = Paragraph(child, doc)
            style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
            original_text = str(getattr(p, "text", "") or "")
            kind = _paragraph_kind(style_name, original_text)
            target_len = estimate_text_length(original_text)
            if kind == "blank":
                min_len = 0
                max_len = 0
            elif kind == "heading":
                min_len = max(4, int(target_len * 0.5))
                max_len = max(12, int(target_len * 1.2))
            else:
                min_len = max(20, int(target_len * 0.7))
                max_len = max(60, int(target_len * 1.3))
            out.append(
                TemplateParagraph(
                    index=idx,
                    style=style_name,
                    kind=kind,
                    original_text=original_text,
                    target_len=target_len,
                    min_len=min_len,
                    max_len=max_len,
                )
            )
            idx += 1
        elif tag.endswith("}tbl"):
            from docx.table import Table  # type: ignore
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in list(getattr(cell, "paragraphs", []) or []):
                        style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
                        original_text = str(getattr(p, "text", "") or "")
                        kind = "table_cell" if original_text.strip() else "blank"
                        target_len = estimate_text_length(original_text)
                        if kind == "blank":
                            min_len = 0
                            max_len = 0
                        elif kind == "heading":
                            min_len = max(4, int(target_len * 0.5))
                            max_len = max(12, int(target_len * 1.2))
                        else:
                            min_len = max(10, int(target_len * 0.7))
                            max_len = max(40, int(target_len * 1.3))
                        out.append(
                            TemplateParagraph(
                                index=idx,
                                style=style_name,
                                kind=kind,
                                original_text=original_text,
                                target_len=target_len,
                                min_len=min_len,
                                max_len=max_len,
                            )
                        )
                        idx += 1
    if len(out) > max_paragraphs:
        raise ValueError(f"文档段落过多（{len(out)}），请先精简后再仿写")
    return out


def _style_sample(template: List[TemplateParagraph], *, max_chars: int = 2000) -> str:
    parts: List[str] = []
    total = 0
    for item in template:
        if item.kind not in {"body", "heading"}:
            continue
        s = (item.original_text or "").strip()
        if not s:
            continue
        if total + len(s) > max_chars:
            remain = max_chars - total
            if remain <= 0:
                break
            s = s[:remain]
        parts.append(s)
        total += len(s)
        if total >= max_chars:
            break
    return "\n\n".join(parts).strip()


def _template_language(template: List[TemplateParagraph]) -> str:
    sample = _style_sample(template, max_chars=3000)
    if sample:
        return detect_language_hint(sample)
    all_text = "\n\n".join((p.original_text or "").strip() for p in template if (p.original_text or "").strip())
    all_text = all_text[:3000]
    return detect_language_hint(all_text)


def _build_prompt(
    *,
    instruction_text: str,
    template: List[TemplateParagraph],
    style_sample: str,
) -> tuple[str, str]:
    spec = []
    for item in template:
        if item.kind == "blank":
            continue
        spec.append(
            {
                "paragraph_index": item.index,
                "kind": item.kind,
                "style": item.style,
                "min_len": item.min_len,
                "max_len": item.max_len,
                "original_excerpt": (item.original_text or "").strip()[:160],
                "case_sensitive_terms": _extract_case_sensitive_terms(item.original_text or ""),
            }
        )
    language_hint = detect_language_hint(style_sample) if style_sample else _template_language(template)
    sys_prompt = (
        "你是一个专业写作助手。你的任务是基于参考文档的写作风格与段落结构，生成新的段落文本。\n"
        "你必须只输出严格的 JSON（不要 Markdown，不要解释，不要代码块，不要多余文本）。\n"
        "输出必须为一个 JSON 数组，每个元素是对象：{\"paragraph_index\": number, \"text\": string}。\n"
        "paragraph_index 必须来自输入的段落清单，且每个 paragraph_index 恰好出现一次（覆盖所有非空段落）。\n"
        "输出语言必须严格与参考文档一致（由 style_sample 决定），即使用户输入内容是另一种语言，也必须先翻译再输出到目标语言。\n"
        "如果目标语言为英文：使用自然、专业的英文；专有名词/姓名/邮箱/电话号码保持原样或给出合理英文写法，不要乱改。\n"
        "不得编造事实或经历；遇到信息缺失时用更通用的表述而不是杜撰。\n"
        "对每个段落的 case_sensitive_terms，必须保持完全一致（包含大小写与符号），不得改写，不得改变大小写。\n"
    )
    user_payload = {
        "instruction": (instruction_text or "").strip(),
        "style_sample": style_sample,
        "language_hint": language_hint,
        "paragraphs": spec,
        "rules": {
            "keep_structure": True,
            "respect_length": True,
            "no_markdown": True,
        },
    }
    user_prompt = json.dumps(user_payload, ensure_ascii=False)
    return sys_prompt, user_prompt


def _parse_llm_json(text: str) -> List[Dict[str, Any]]:
    raw = (text or "").strip()
    if not raw:
        raise ValueError("LLM 返回空内容")
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return data
    except Exception:
        pass
    start = raw.find("[")
    end = raw.rfind("]")
    if start >= 0 and end > start:
        snippet = raw[start : end + 1]
        data = json.loads(snippet)
        if isinstance(data, list):
            return data
    raise ValueError("LLM 返回内容不是有效 JSON 数组")


def generate_rewrite_mapping(
    *,
    llm: ChatOpenAI,
    instruction_text: str,
    template: List[TemplateParagraph],
) -> Dict[int, str]:
    style_sample = _style_sample(template)
    sys_prompt, user_prompt = _build_prompt(
        instruction_text=instruction_text,
        template=template,
        style_sample=style_sample,
    )
    msg = llm.invoke([
        ("system", sys_prompt),
        ("user", user_prompt),
    ])
    content = (getattr(msg, "content", "") or "").strip()
    items = _parse_llm_json(content)
    mapping: Dict[int, str] = {}
    for obj in items:
        if not isinstance(obj, dict):
            continue
        idx = obj.get("paragraph_index")
        txt = obj.get("text")
        if not isinstance(idx, int):
            continue
        if not isinstance(txt, str):
            continue
        mapping[idx] = txt

    required = [p.index for p in template if p.kind != "blank"]
    for idx in required:
        if idx not in mapping:
            mapping[idx] = ""
    return mapping


def _extract_case_sensitive_terms(text: str, *, max_terms: int = 20) -> List[str]:
    s = (text or "").strip()
    if not s:
        return []
    patterns = [
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        r"\b[A-Z][A-Z0-9]{1,}\b",
        r"\b[A-Z][a-z]+[A-Z][A-Za-z0-9]*\b",
        r"C\+\+|\.NET",
    ]
    found: List[str] = []
    for pat in patterns:
        for m in re.finditer(pat, s):
            tok = (m.group(0) or "").strip()
            if not tok:
                continue
            if tok in found:
                continue
            found.append(tok)
            if len(found) >= max_terms:
                return found
    return found


def _enforce_case_terms(reference_text: str, new_text: str) -> str:
    out = new_text or ""
    terms = _extract_case_sensitive_terms(reference_text or "")
    if not terms:
        return out
    for term in terms:
        if not term:
            continue
        if re.fullmatch(r"[A-Za-z0-9]+", term or ""):
            pat = re.compile(rf"\b{re.escape(term)}\b", flags=re.IGNORECASE)
        else:
            pat = re.compile(re.escape(term), flags=re.IGNORECASE)
        out = pat.sub(term, out)
    return out


def _clear_paragraph_content(paragraph: Any) -> None:
    p_elm = getattr(paragraph, "_p", None)
    if p_elm is None:
        return
    children = list(p_elm)
    for child in children:
        if str(getattr(child, "tag", "")).endswith("}pPr"):
            continue
        try:
            p_elm.remove(child)
        except Exception:
            pass


def _clone_run_properties(src_run: Any, dst_run: Any) -> None:
    if src_run is None or dst_run is None:
        return
    try:
        src_style = getattr(src_run, "style", None)
        if src_style is not None:
            dst_run.style = src_style
    except Exception:
        pass

    try:
        src_rPr = getattr(getattr(src_run, "_element", None), "rPr", None)
        if src_rPr is None:
            return
        dst_elm = getattr(dst_run, "_element", None)
        if dst_elm is None:
            return
        dst_rPr = dst_elm.get_or_add_rPr()
        for child in list(dst_rPr):
            dst_rPr.remove(child)
        src_copy = deepcopy(src_rPr)
        for k, v in getattr(src_copy, "attrib", {}).items():
            dst_rPr.set(k, v)
        for child in list(src_copy):
            dst_rPr.append(child)
    except Exception:
        pass


def render_rewritten_docx(reference_docx_bytes: bytes, mapping: Dict[int, str]) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("缺少依赖：python-docx") from e

    doc = Document(BytesIO(reference_docx_bytes))
    body = getattr(getattr(doc, "_element", None), "body", None)
    if body is None:
        paras = list(getattr(doc, "paragraphs", []) or [])
        for idx, p in enumerate(paras):
            if idx not in mapping:
                continue
            original_text = str(getattr(p, "text", "") or "")
            new_text = _enforce_case_terms(original_text, mapping[idx])
            ref_run = None
            try:
                runs = list(getattr(p, "runs", []) or [])
                if runs:
                    ref_run = runs[0]
            except Exception:
                ref_run = None
            _clear_paragraph_content(p)
            new_run = p.add_run(new_text)
            _clone_run_properties(ref_run, new_run)
    else:
        i = 0
        for child in list(body):
            tag = str(getattr(child, "tag", ""))
            if tag.endswith("}p"):
                from docx.text.paragraph import Paragraph  # type: ignore
                p = Paragraph(child, doc)
                if i in mapping:
                    original_text = str(getattr(p, "text", "") or "")
                    new_text = _enforce_case_terms(original_text, mapping[i])
                    ref_run = None
                    try:
                        runs = list(getattr(p, "runs", []) or [])
                        if runs:
                            ref_run = runs[0]
                    except Exception:
                        ref_run = None
                    _clear_paragraph_content(p)
                    new_run = p.add_run(new_text)
                    _clone_run_properties(ref_run, new_run)
                i += 1
            elif tag.endswith("}tbl"):
                from docx.table import Table  # type: ignore
                tbl = Table(child, doc)
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in list(getattr(cell, "paragraphs", []) or []):
                            if i in mapping:
                                original_text = str(getattr(p, "text", "") or "")
                                new_text = _enforce_case_terms(original_text, mapping[i])
                                ref_run = None
                                try:
                                    runs = list(getattr(p, "runs", []) or [])
                                    if runs:
                                        ref_run = runs[0]
                                except Exception:
                                    ref_run = None
                                _clear_paragraph_content(p)
                                new_run = p.add_run(new_text)
                                _clone_run_properties(ref_run, new_run)
                            i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def rewrite_docx_mvp(
    *,
    reference_docx_bytes: bytes,
    instruction_text: str,
    llm_config: LlmConfig,
    max_paragraphs: int = 400,
) -> bytes:
    template = extract_docx_template(reference_docx_bytes, max_paragraphs=max_paragraphs)
    llm = ChatOpenAI(
        temperature=0,
        max_retries=3,
        base_url=llm_config.base_url,
        model=llm_config.model,
        api_key=llm_config.api_key,
    )
    mapping = generate_rewrite_mapping(
        llm=llm,
        instruction_text=instruction_text,
        template=template,
    )
    return render_rewritten_docx(reference_docx_bytes, mapping)
